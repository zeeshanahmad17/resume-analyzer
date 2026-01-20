import pdfplumber
import re
from typing import Optional
from docx import Document
from config import Config

class DocumentParser:
    """Extract and clean text from resumes"""

    def extract_content(self, file_path: str) -> str:
        """
        Extract text from a file.

        Args:
            file_path: Path to the file.

        Returns:
            str: Extracted text from the file.
        """
        # Determine file type
        file_extension = file_path.lower().split('.')[-1]

        if file_extension == 'pdf':
            return self._extract_pdf_content(file_path)
        elif file_extension == 'txt':
            return self._extract_txt_content(file_path)
        elif file_extension in ['doc', 'docx']:
            return self._extract_docx_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF"""
        text_content = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            full_text = "\n".join(text_content)
            return self._clean_text(full_text)

        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")

    def _extract_txt_content(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                return self._clean_text(text)
        except Exception as e:
            raise ValueError(f"Error extracting text from TXT: {str(e)}")
        
    def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text_content = [paragraph.text for paragraph in doc.paragraphs]
            full_text = "\n".join(text_content)
            return self._clean_text(full_text)
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Remove noise and normalize whitespace."""

        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep essential punctuation
        text = re.sub(r'[^\w\s.,;:()\-@/#&+]', '', text)

        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)

        return text.strip()

    def validate_file(self, file_path: str) -> bool:
        """Valide that file is a readable file."""
        file_extension = file_path.lower().split('.')[-1]

        try:
            if file_extension == 'pdf':
                with pdfplumber.open(file_path) as pdf:
                    return len(pdf.pages) > 0
            elif file_extension == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return len(f.read()) > 0
            elif file_extension in ['doc', 'docx']:
                doc = Document(file_path)
                return len(doc.paragraphs) > 0
            return False
        except:
            return False